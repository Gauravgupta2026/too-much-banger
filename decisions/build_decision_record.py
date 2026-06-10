from __future__ import annotations

import csv
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
DOCX_PATH = BASE_DIR / "too-much-banger-decision-tracker.docx"
CSV_PATH = BASE_DIR / "too-much-banger-decision-tracker.csv"


@dataclass
class Decision:
    decision_id: str
    title: str
    phase: str
    stage: str
    category: str
    date_decided: str
    status: str
    owner: str
    approver: str
    summary: str
    primary_driver: str
    facts: list[str]
    assumptions: list[str]
    chosen_option: str
    rejected_options: list[str]
    tradeoff: str
    risk_accepted: str
    confidence: str
    revisit_trigger: str
    revisit_date: str
    outcome_status: str
    context: str
    decision_statement: str
    why_now: str


DECISIONS = [
    Decision(
        decision_id="TMB-001",
        title="We are building a public archive about the Bangerlore discourse, not the party homepage itself.",
        phase="Ideation",
        stage="Discovery",
        category="Scope",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The product is a public-facing archive of the online reaction, with enough context for outsiders to understand the incident before entering the archive wall.",
        primary_driver="The site needed a clear premise and audience before design or implementation decisions would make sense.",
        facts=[
            "The prompt asked for an archival website on toomuchbanger.in, not a Bangerlore marketing site.",
            "The user wanted a brief explanation for visitors who know nothing about the event.",
            "The rough FigJam flow explicitly placed context before the 'Meanwhile on the internet' section."
        ],
        assumptions=[
            "A public audience will need narrative onboarding before seeing screenshots or quote cards.",
            "The archive is more interesting if it frames the discourse rather than defending the party."
        ],
        chosen_option="Frame the product as a public archive with a context-first opening, then a satirical transition into the internet wall.",
        rejected_options=[
            "Build a generic event microsite with a small reactions section.",
            "Start immediately with tweets and assume visitors already know the lore."
        ],
        tradeoff="We spend screen real estate on context instead of maximizing the number of posts above the fold.",
        risk_accepted="People already deep in the lore may find the opening slower than necessary.",
        confidence="High",
        revisit_trigger="If users consistently skip the opening and only engage with the archive wall.",
        revisit_date="After first public traffic review",
        outcome_status="Too Early",
        context="The concept only becomes legible if the site explains why the archive exists and what cultural argument it is preserving.",
        decision_statement="The first part of the site will explain what Bangerlore was, what happened the next day, and why this archive exists before it shows the wall of posts.",
        why_now="Without this call, design would drift between party promo, satire site, and screenshot dump."
    ),
    Decision(
        decision_id="TMB-002",
        title="We are using playful satire instead of a hostile wall-of-shame tone.",
        phase="Ideation",
        stage="Discovery",
        category="UX",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The site voice is intentionally comic and absurd, with enough distance to stay public-safe and not collapse into mean-spirited callout culture.",
        primary_driver="Tone choice materially changes the legal, reputational, and design posture of the project.",
        facts=[
            "The selected planning input explicitly chose 'Playful satire' over sharper roast or museum archive.",
            "The design note leaned toward 'morning after house party' and 'walk of shame' motifs, but not direct harassment."
        ],
        assumptions=[
            "The archive will travel more safely if the humor is about the discourse rather than about humiliating individual people."
        ],
        chosen_option="Use playful satire as the default voice across copy, layout, reactions, and section names.",
        rejected_options=[
            "Run a harsher roast-first version with more aggressive copy.",
            "Make it read like a neutral museum archive."
        ],
        tradeoff="The site gives up some shock value in exchange for broader shareability and lower moderation risk.",
        risk_accepted="Some visitors may want sharper commentary than the product is willing to provide.",
        confidence="High",
        revisit_trigger="If real-content review shows the humor still feels too adversarial or too flat.",
        revisit_date="Before swapping mock posts for real posts",
        outcome_status="Too Early",
        context="Tone controls what kinds of interactions feel acceptable and what design language will still feel intentional instead of cruel.",
        decision_statement="The archive will make fun of the situation and the discourse mechanics, not turn into a public humiliation board.",
        why_now="Tone had to be locked before naming reactions, writing copy, or choosing the visual system."
    ),
    Decision(
        decision_id="TMB-003",
        title="We are shipping V1 with mock posts and no live X embeds.",
        phase="Ideation",
        stage="Scoping",
        category="Scope",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The first version uses structured mock posts so the design can be built and reviewed without depending on live embeds, scraping, or final curation.",
        primary_driver="Design iteration needed to move before final content curation and platform integration were ready.",
        facts=[
            "The request explicitly said to use mock tweets for now and real ones later.",
            "The plan ruled out live X embeds in V1 because they reduce layout control and add instability."
        ],
        assumptions=[
            "The design risks are higher than the content ingestion risks in the first pass.",
            "A clean schema for mock data will make the real-content swap straightforward later."
        ],
        chosen_option="Use a local typed mock-post dataset and render custom archive cards rather than live embedded posts.",
        rejected_options=[
            "Embed real X posts in V1.",
            "Delay the whole build until the real post set is finalized."
        ],
        tradeoff="The archive is visually stronger now but not yet source-authentic.",
        risk_accepted="People reviewing the first version may over-index on placeholder language or counts.",
        confidence="High",
        revisit_trigger="When the real post set is approved and stable enough to replace the mocks.",
        revisit_date="Before launch on the production domain",
        outcome_status="Working",
        context="The product needed a buildable content model immediately, and live social embeds would have pushed decisions into third-party constraints too early.",
        decision_statement="Mock posts are the source of truth for V1 design and interaction work; real posts will be a later data swap, not a redesign.",
        why_now="The site could not move into implementation until the content dependency was temporarily removed."
    ),
    Decision(
        decision_id="TMB-004",
        title="We are building this as a new standalone Next.js app instead of extending the existing Python repo.",
        phase="Development",
        stage="Delivery",
        category="Technical",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The website lives in its own `too-much-banger` app directory, separate from the existing Twitter-agent codebase, to avoid coupling an unrelated frontend onto a Python agent project.",
        primary_driver="The existing workspace root was a Python Twitter-agent project and did not contain a usable web scaffold.",
        facts=[
            "Repo inspection found a Python app, not an existing web frontend to extend.",
            "No nearby Next.js/Vite/Astro scaffold was present in the explored tree."
        ],
        assumptions=[
            "The site will evolve faster if its frontend stack is isolated from the existing agent code."
        ],
        chosen_option="Create a standalone Next.js 15 app under `too-much-banger` and leave the Python project untouched.",
        rejected_options=[
            "Bolt a frontend onto the Python agent repo root.",
            "Start a fully separate repository before validating the product."
        ],
        tradeoff="The workspace now contains two unrelated products, but the code boundaries are explicit.",
        risk_accepted="Deployment and project-level tooling need a little extra care because the web app is a nested project.",
        confidence="High",
        revisit_trigger="If the site graduates into its own long-lived product with a separate team or release process.",
        revisit_date="When the project is live and stable",
        outcome_status="Working",
        context="This was an implementation hygiene decision as much as a product one. The wrong choice here would have created avoidable technical debt on day one.",
        decision_statement="The site is a separate application with its own package.json, build, deployment config, and Supabase integration.",
        why_now="Stack choice affects every file and deployment decision that followed."
    ),
    Decision(
        decision_id="TMB-005",
        title="We are using Vercel for hosting and keeping the deployment path aligned with the custom domain from day one.",
        phase="Development",
        stage="Launch",
        category="Rollout",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The site is configured for Vercel deployment with the expectation that `toomuchbanger.in` will be attached there once auth, env vars, and DNS are ready.",
        primary_driver="The project needed a fast deployment path with first-class support for a modern React frontend and lightweight APIs.",
        facts=[
            "The planning choice explicitly selected Vercel as the recommended hosting platform.",
            "The app now includes `vercel.json` and passes a production Next.js build."
        ],
        assumptions=[
            "The user controls DNS for `toomuchbanger.in`.",
            "Serverless routes are sufficient for reactions and interest capture at this scale."
        ],
        chosen_option="Target Vercel as the primary host and wire the app to that deployment model from the start.",
        rejected_options=[
            "Keep the host undecided until after implementation.",
            "Optimize first for Cloudflare Pages or a custom server."
        ],
        tradeoff="The deployment story is simpler, but it depends on Vercel account auth and environment management.",
        risk_accepted="The actual launch is blocked until Vercel authentication and DNS changes are completed.",
        confidence="Medium",
        revisit_trigger="If account constraints, pricing, or domain setup make Vercel materially worse than expected.",
        revisit_date="Before production launch",
        outcome_status="Holding",
        context="The domain requirement was part of the original ask, so hosting could not remain vague.",
        decision_statement="Vercel is the deployment target and the app is shaped to deploy there cleanly.",
        why_now="Deployment assumptions affect project structure, environment naming, and API design."
    ),
    Decision(
        decision_id="TMB-006",
        title="We are using Supabase for persistent reactions and interest submissions, with graceful fallback when env vars are missing.",
        phase="Development",
        stage="Delivery",
        category="Technical",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The product supports shared reaction counts and CTA submissions through Supabase, but the UI is intentionally resilient when the backend is not configured yet.",
        primary_driver="The product needed persistent interaction data without forcing a heavy backend build.",
        facts=[
            "The planning choice explicitly selected persistent counts rather than local-only reactions.",
            "The app now includes Supabase-ready API routes, schema SQL, and an env template.",
            "Browser tests confirmed reactions and form submissions fail gracefully without configured env vars."
        ],
        assumptions=[
            "Supabase is enough for this interaction model and does not require a custom backend in V1."
        ],
        chosen_option="Use Supabase as the persistence layer and design the frontend to degrade cleanly when the service is not yet connected.",
        rejected_options=[
            "Ship local-only reactions with no backend path.",
            "Build a custom database/API layer before validating the site."
        ],
        tradeoff="The system is operationally lightweight, but launch readiness depends on external configuration that is not done yet.",
        risk_accepted="Until env vars and schema are live, the site behaves like a visual prototype for those features.",
        confidence="High",
        revisit_trigger="If reaction volume, moderation needs, or analytics requirements exceed what this simple schema can support.",
        revisit_date="After the first real traffic cycle",
        outcome_status="Working",
        context="Persistence mattered to the concept, but backend complexity had to stay proportional to the scope.",
        decision_statement="Supabase is the data layer for global reactions and 'want in' submissions, and the app shows explicit fallback copy when it is absent.",
        why_now="Interaction design and API design had to converge before implementation was complete."
    ),
    Decision(
        decision_id="TMB-007",
        title="We are committing to a morning-after scrapbook visual system with context, comic transition, masonry cards, and evidence-board leaderboard.",
        phase="Development",
        stage="Delivery",
        category="UX",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The design direction is not a safe template. It deliberately uses a dark poster-like hero, comic-panel transition, scrapbook card wall, and playful evidence-board language.",
        primary_driver="The product would fail if the design looked like a generic startup landing page with tweets pasted into it.",
        facts=[
            "The reference notes called for a 'morning after house party' layout and absurd archive motifs.",
            "The rough sketches proposed the 'Meanwhile on the internet' transition and vote-on-hover behavior."
        ],
        assumptions=[
            "A more theatrical layout is an asset here because the subject is cultural spectacle, not SaaS conversion."
        ],
        chosen_option="Use a bold, editorial visual system that feels like a messy but curated internet artifact.",
        rejected_options=[
            "Build a clean minimal landing page with standard testimonial-card treatment.",
            "Mirror the exact Bangerlore aesthetic too closely."
        ],
        tradeoff="The design is less neutral and may polarize, but it is far more memorable and aligned with the premise.",
        risk_accepted="Some viewers may read the style as unserious even though the product is intentionally structured.",
        confidence="High",
        revisit_trigger="If real-content review shows the styling overwhelms legibility or makes the archive hard to scan.",
        revisit_date="Before final content swap",
        outcome_status="Working",
        context="This is a design-led project. The visual system is part of the product logic, not cosmetic frosting.",
        decision_statement="The site will behave like a curated aftermath archive, not a standard marketing page with social content blocks.",
        why_now="Visual direction needed to be locked before component structure and copy could settle."
    ),
    Decision(
        decision_id="TMB-008",
        title="We are treating launch readiness as 'code complete, ops pending' rather than pretending the site is already live.",
        phase="Deployment",
        stage="Launch",
        category="Operational",
        date_decided="2026-06-10",
        status="Decided",
        owner="Gaurav Gupta",
        approver="Gaurav Gupta",
        summary="The build is verified locally and browser-tested, but the team is explicitly tracking the remaining operational blockers instead of masking them.",
        primary_driver="There was a risk of confusing implementation completeness with actual launch completeness.",
        facts=[
            "Lint, typecheck, build, and local production-browser smoke tests passed.",
            "Vercel CLI was installed successfully but was not authenticated on this machine.",
            "Supabase project values and production DNS changes were not yet applied."
        ],
        assumptions=[
            "It is better to record the remaining gaps explicitly than to declare the project effectively done."
        ],
        chosen_option="Mark the site as deployment-ready in code but blocked operationally by account auth, environment setup, and DNS.",
        rejected_options=[
            "Describe the site as fully launched even though auth and env setup were unfinished.",
            "Delay all decision recording until after launch."
        ],
        tradeoff="The record is more honest, but it also documents that launch work is still outstanding.",
        risk_accepted="The project could stall if the operational handoff is not finished promptly.",
        confidence="High",
        revisit_trigger="When Vercel auth is completed, Supabase is configured, and the domain resolves to the deployed app.",
        revisit_date="At deployment completion",
        outcome_status="Holding",
        context="This is the difference between a finished codebase and a finished product.",
        decision_statement="We will track the site as build-complete but not launch-complete until hosting, backend config, and domain wiring are done.",
        why_now="The decision log should reflect reality, not just implementation momentum."
    ),
]


def set_run_font(run, name: str, size: int | None = None, bold: bool | None = None, color: str | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line: float | None = None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line is not None:
        fmt.line_spacing = line


def add_bottom_border(paragraph, color: str = "DADCE0", size: str = "6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def ensure_styles(document: Document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, "000000"),
        ("Heading 2", 16, 18, 6, "000000"),
        ("Heading 3", 14, 16, 4, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15

    if "Decision Meta" not in styles:
        meta = styles.add_style("Decision Meta", WD_STYLE_TYPE.PARAGRAPH)
        meta.font.name = "Arial"
        meta._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        meta._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        meta.font.size = Pt(9)
        meta.font.color.rgb = RGBColor.from_string("555555")
        meta.paragraph_format.space_before = Pt(0)
        meta.paragraph_format.space_after = Pt(6)
        meta.paragraph_format.line_spacing = 1.0


def add_bullets(document: Document, items: list[str]):
    for item in items:
        paragraph = document.add_paragraph(style="Normal")
        paragraph.style = document.styles["Normal"]
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        set_paragraph_spacing(paragraph, 0, 4, 1.15)
        run = paragraph.add_run(f"• {item}")
        set_run_font(run, "Arial", 11)


def add_label_paragraph(document: Document, label: str, text: str):
    paragraph = document.add_paragraph(style="Normal")
    set_paragraph_spacing(paragraph, 0, 8, 1.15)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, "Arial", 11, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, "Arial", 11)


def build_docx():
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    ensure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, 0, 3, 1.0)
    title_run = title.add_run("Too Much Banger - Product Decision Tracker")
    set_run_font(title_run, "Arial", 26, bold=False, color="000000")

    subtitle = document.add_paragraph(style="Decision Meta")
    subtitle_run = subtitle.add_run("Project: Too Much Banger | Last updated: 2026-06-10 | Owner: Gaurav Gupta")
    set_run_font(subtitle_run, "Arial", 9, color="555555")
    add_bottom_border(subtitle)

    intro_paragraphs = [
        "This document tracks the main product and delivery decisions behind Too Much Banger. It exists so the reasoning does not get rewritten after the fact once the site is live.",
        "The scope here is the first public version: concept, tone, build choices, interaction model, and launch posture. The goal is not to make the story look elegant. The goal is to make the calls defensible later.",
        "Some of these decisions are settled. Some are intentionally temporary and will be revisited when real posts, Supabase configuration, and the production domain are in place."
    ]
    for paragraph_text in intro_paragraphs:
        paragraph = document.add_paragraph(paragraph_text, style="Normal")
        set_paragraph_spacing(paragraph, 0, 8, 1.15)

    phases = ["Ideation", "Development", "Deployment"]
    for phase in phases:
        document.add_paragraph(phase, style="Heading 1")
        phase_decisions = [decision for decision in DECISIONS if decision.phase == phase]
        for decision in phase_decisions:
            heading = document.add_paragraph(decision.title, style="Heading 2")
            set_paragraph_spacing(heading, 18, 6, 1.15)

            meta = document.add_paragraph(style="Decision Meta")
            meta_run = meta.add_run(
                f"{decision.decision_id} | Stage: {decision.stage} | Category: {decision.category} | Status: {decision.status} | Confidence: {decision.confidence}"
            )
            set_run_font(meta_run, "Arial", 9, color="555555")

            add_label_paragraph(document, "Context", decision.context)
            add_label_paragraph(document, "Decision", decision.decision_statement)
            add_label_paragraph(document, "Why now", decision.why_now)

            facts_heading = document.add_paragraph(style="Heading 3")
            facts_heading.add_run("Facts")
            add_bullets(document, decision.facts)

            if decision.assumptions:
                assumptions_heading = document.add_paragraph(style="Heading 3")
                assumptions_heading.add_run("Assumptions")
                add_bullets(document, decision.assumptions)

            options_heading = document.add_paragraph(style="Heading 3")
            options_heading.add_run("Options considered")
            chosen = document.add_paragraph(style="Normal")
            chosen.paragraph_format.left_indent = Inches(0.25)
            chosen.paragraph_format.first_line_indent = Inches(-0.25)
            set_paragraph_spacing(chosen, 0, 4, 1.15)
            chosen_label = chosen.add_run("• Chosen: ")
            set_run_font(chosen_label, "Arial", 11, bold=True)
            chosen_text = chosen.add_run(decision.chosen_option)
            set_run_font(chosen_text, "Arial", 11)
            for rejected in decision.rejected_options:
                rejected_para = document.add_paragraph(style="Normal")
                rejected_para.paragraph_format.left_indent = Inches(0.25)
                rejected_para.paragraph_format.first_line_indent = Inches(-0.25)
                set_paragraph_spacing(rejected_para, 0, 4, 1.15)
                rejected_label = rejected_para.add_run("• Rejected: ")
                set_run_font(rejected_label, "Arial", 11, bold=True)
                rejected_text = rejected_para.add_run(rejected)
                set_run_font(rejected_text, "Arial", 11)

            add_label_paragraph(document, "Tradeoff accepted", decision.tradeoff)
            add_label_paragraph(document, "Risk accepted", decision.risk_accepted)
            add_label_paragraph(
                document,
                "Revisit trigger",
                f"{decision.revisit_trigger}. Revisit date: {decision.revisit_date}."
            )
            add_label_paragraph(
                document,
                "Owner and approver",
                f"Owner: {decision.owner}. Approver: {decision.approver}. Key contributor: Codex."
            )

    revision_heading = document.add_paragraph("Revision log", style="Heading 1")
    set_paragraph_spacing(revision_heading, 20, 6, 1.15)
    revision_entry = document.add_paragraph(style="Normal")
    set_paragraph_spacing(revision_entry, 0, 8, 1.15)
    revision_bold = revision_entry.add_run("2026-06-10: ")
    set_run_font(revision_bold, "Arial", 11, bold=True)
    revision_text = revision_entry.add_run(
        "Initial record created after planning, implementation, local verification, and deployment preparation."
    )
    set_run_font(revision_text, "Arial", 11)

    document.save(DOCX_PATH)


def build_csv():
    rows = []
    for decision in DECISIONS:
        rows.append(
            {
                "Decision ID": decision.decision_id,
                "Decision Title": decision.title,
                "Product": "Too Much Banger",
                "Stage": decision.stage,
                "Category": decision.category,
                "Date Decided": decision.date_decided,
                "Status": decision.status,
                "Owner": decision.owner,
                "Approver": decision.approver,
                "Decision Summary": decision.summary,
                "Primary Driver": decision.primary_driver,
                "Facts Snapshot": " | ".join(decision.facts),
                "Key Assumptions": " | ".join(decision.assumptions),
                "Chosen Option": decision.chosen_option,
                "Rejected Options": " | ".join(decision.rejected_options),
                "Main Tradeoff": decision.tradeoff,
                "Risk Accepted": decision.risk_accepted,
                "Confidence": decision.confidence,
                "Revisit Trigger": decision.revisit_trigger,
                "Revisit Date": decision.revisit_date,
                "Outcome Status": decision.outcome_status,
                "Doc Link": DOCX_PATH.name,
                "Notion Link": "",
                "Notes": "",
            }
        )

    with CSV_PATH.open("w", newline="", encoding="utf-8") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def main():
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_csv()


if __name__ == "__main__":
    main()
